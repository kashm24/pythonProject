import xlrd2
from tkinter import *
from tkinter import messagebox
from datetime import datetime
from tkinter import ttk
import tkinter as tk
import requests
import docx
from tkinter.ttk import Combobox
import matplotlib.pyplot as plt
from docx import document
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from tkcalendar import Calendar

def donwloade(event): #загрузка базы
    try:
        files = open('vullist.xlsx', "wb")
        url = 'https://bdu.fstec.ru/files/documents/vullist.xlsx'
        headers = {
        'User-Agent': 'My User Agent 1.0',
        }
        response = requests.get(url, headers=headers)
        files.write(response.content)
        files.close()
    except requests.exceptions.ConnectionError:
        messagebox.showerror('Ошибка',
                             'Неудалось загрузить файл, проверьте интернет подключение')
def by_date(event):  #Фильтрация по дате
    labelDateInfo.configure(state=NORMAL)
    textBoxFromDate.configure(state=NORMAL)
    textBoxToDate.configure(state=NORMAL)
    labelFromDate.configure(state=NORMAL)
    labelToDate.configure(state=NORMAL)
def without_date(event):  #Фильтрация без даты
    textBoxFromDate.delete(0, END)
    textBoxToDate.delete(0, END)
    labelDateInfo.configure(state=DISABLED)
    textBoxFromDate.configure(state=DISABLED)
    textBoxToDate.configure(state=DISABLED)
    labelFromDate.configure(state=DISABLED)
    labelToDate.configure(state=DISABLED)
def AnalysysWithDate(event):  # Функция для проверки правильности ввода даты
    chrb = radioButtonDateVar.get()  # состояние кнопок, 0 - за все время, 1 - по дате
    if chrb == 0:
        Analysys(event)
    else:
        dataFrom = textBoxFromDate.get()
        dataTo = textBoxToDate.get()
        if len(dataFrom and dataTo) == 10 and (dataFrom[2] and dataTo[2]) == '.' and (dataFrom[5] and dataTo[5]) == '.' and dataFrom[6:].isnumeric() and dataTo[6:].isnumeric() and dataFrom[:2].isnumeric() and dataTo[:2].isnumeric() and dataFrom[3:5].isnumeric() and dataTo[3:5].isnumeric():
          tsFrom = datetime(year=int(dataFrom[6:]), month=int(dataFrom[3:5]), day=int(dataFrom[:2]))
          tsTo = datetime(year=int(dataFrom[6:]), month=int(dataFrom[3:5]), day=int(dataFrom[:2]))
          if (tsFrom.date and tsTo.day) > 0 and (tsFrom.day and tsTo.day) < 32 and (tsFrom.month and tsTo.month) > 0 and (tsFrom.month and tsTo.month) < 13 and (tsFrom.year and tsTo.year) > 1900:
            dataFrom = datetime.strptime(textBoxFromDate.get(), "%d.%m.%Y")
            dataTo = datetime.strptime(textBoxToDate.get(), "%d.%m.%Y")
            Analysys(event)
          else:
            messagebox.showerror('Ошибка',
                                 'Некорректно введена дата')
        else:
          messagebox.showerror('Ошибка',
                               'Некорректно введена дата')

def Analysys(event):  # Функция поиска уязвимостей
    try:
        try:
            workbook = xlrd2.open_workbook('vullist.xlsx')
            sheet = workbook.sheet_by_index(0)  # получаем доступ к первой странице
            row = sheet.nrows  # определяем количество записей
            if row != 0:
                names = sheet.col_values(4)
                status = sheet.col_values(14)
                danger_lavels = sheet.col_values(12)
                chrb = radioButtonDateVar.get()
                ddd = sheet.col_values(9)
                DDD = var1.get()
                sow=combo.get()
                global danger_low, danger_middle, danger_hight, danger_crit # уровни уязвимости
                danger_crit, danger_hight, danger_middle, danger_low = 0, 0, 0, 0
                if chrb == 0: # если 0 то без даты фильтровать
                    dataFrom = datetime.strptime('01.01.2001',"%d.%m.%Y")
                    dataTo = datetime.strptime('17.06.2023', "%d.%m.%Y")
                else: # тогда
                    dataFrom = datetime.strptime(textBoxFromDate.get(), "%d.%m.%Y")
                    dataTo = datetime.strptime(textBoxToDate.get(), "%d.%m.%Y")

                for i in range(9, row):
                    if ddd[i] != '':
                        ddd[i] = datetime.strptime(ddd[i], "%d.%m.%Y")
                    else:
                        ddd[i] = datetime.strptime('01.01.2001', "%d.%m.%Y")
                for i in range(4, row):
                    if (str(ddd[i]) >= str(dataFrom)) and (str(ddd[i]) <= str(dataTo)):
                        if DDD == 1:
                            if status[i].find("Потенциальная уязвимость") >= 0 and names[i].find(sow) >= 0:
                                if danger_lavels[i][0] == 'К':  # Критический
                                    danger_crit += 1

                                elif danger_lavels[i][0] == 'В':  # Высокий
                                    danger_hight += 1

                                elif danger_lavels[i][0] == 'С':  # Средний
                                    danger_middle += 1

                                else:  # Низкий
                                    danger_low += 1
                        else:
                            if names[i].find(sow) >= 0:

                              if danger_lavels[i][0] == 'К':  # Критический
                                danger_crit += 1

                              elif danger_lavels[i][0] == 'В':  # Высокий
                                danger_hight += 1

                              elif danger_lavels[i][0] == 'С':  # Средний
                                danger_middle += 1

                              else:  # Низкий
                                danger_low += 1

                labelLowOut['text'] = danger_low
                labelMidOut['text'] = danger_middle
                labelHighOut['text'] = danger_hight
                labelCritOut['text'] = danger_crit
            else:  # сообщения об ошибке с базой
                messagebox.showerror('Ошибка',
                                     'Для анализа необходимо обновить базу')
        except FileNotFoundError:
            messagebox.showerror('Ошибка',
                                 'Для анализа необходимо загрузить (Обновить) базу')
    except xlrd2.biffh.XLRDError:
        messagebox.showerror('Ошибка',
                             'Для анализа необходимо загрузить (Обновить) базу')

def diagramma(event):
    try:
        if danger_low == 0 and danger_middle == 0 and danger_hight == 0 and danger_crit ==0:
            messagebox.showerror('Ошибка',
                                 'Для вывода диаграммы необходимо хоть одно не равное 0')
        else:
            labels = 'Низкий', 'Средний', 'Высокий', 'Критический' #название паев
            sizes = [danger_low, danger_middle, danger_hight, danger_crit] #паи пирога
            colors = ("#c0c0c0", "#fff44f", "#FFAF18",
            "#E52E2A") # цвет паев
            figure, ax1 = plt.subplots()
            explode = (0, 0.1, 0, 0)

            ax1.pie(sizes, wedgeprops=dict(width=1), colors=colors, explode=explode, labels=labels,autopct='%1.1f%%', startangle=90) # пирог диаграммы
            patches, texts, auto = ax1.pie(sizes, wedgeprops=dict(width=1), colors=colors, startangle=90, explode=explode, autopct='%1.1f%%' )
            plt.legend(patches, labels, loc='best', bbox_to_anchor=(0.7, 0.67, 0.5, 0.5)) # расположение легенды в диаграмме

            okno=Tk() # окно с диаграммой
            okno.title("Диаграмма") #название окна
            okno.configure(background='#082567') #цвет фона
            canvas = FigureCanvasTkAgg(figure, master=okno) #поля для пирога
            canvas.get_tk_widget().pack()
            canvas.draw()
    except NameError: #ошибка если запускать без анализа
        messagebox.showerror('Ошибка',
                             'Для вывода диаграммы необходимо провести анализ')

def Clear(event): #очистка полей
    labelLowOut['text'] = "" #оставляет пустым поле низкий
    labelMidOut['text'] = "" #оставляет пустым поле средний
    labelHighOut['text'] = "" #оставляет пустым поле высокий
    labelCritOut['text'] = "" #оставляет пустым поле критический
    textBoxFromDate.delete(0, END) #чистит поле даты
    textBoxToDate.delete(0, END) #чистит поле даты

def SaveDocx(event):  # Функция для сохранения результатов в docx
  if labelLowOut['text'] == "" and labelMidOut['text'] == "" and labelHighOut['text'] == "" and labelCritOut['text'] == "": # если все поля уязвимостей пустые выдть ошибку
      messagebox.showerror('Ошибка',
                           'Для вывода отчёта в документ необходимо провести анализ')
  else: #иначе создать файл ворд
        document = docx.Document()
        document.add_heading(combo.get(), 1) #берем слово по которому искали и добовляем его в заголовок
        document.add_heading('Количество уязвимостей по уровням опасности', 1) # название таблицы
        table = document.add_table(rows=5, cols=2) #размеры таблицы

        nazvcell = table.rows[0].cells # ячейка названия
        nazvcell[0].text = 'Уровень' # 1 столбец название
        nazvcell[1].text = 'Количество' # 2 столбец название

        cellforlow = table.rows[1].cells # ячейка для данных Низкий
        cellforlow[0].text = 'Низкий' # 1 столбец название
        cellforlow[1].text = str(labelLowOut['text']) # 2 столбец количество

        cellformid = table.rows[2].cells #ячейка для данных Средний
        cellformid[0].text = 'Средний' # 1 столбец название
        cellformid[1].text = str(labelMidOut['text']) # 2 столбец количество

        cellsforhigh = table.rows[3].cells #ячейка для данных Высокий
        cellsforhigh[0].text = 'Высокий' # 1 столбец название
        cellsforhigh[1].text = str(labelHighOut['text']) # 2 столбец количество

        cellsforcrit = table.rows[4].cells #ячейка для данных Критический
        cellsforcrit[0].text = 'Критический' # 1 столбец название
        cellsforcrit[1].text = str(labelCritOut['text']) # 2 столбец количество

        document.save('Анализ уязвимостей ' + combo.get() + '.docx') #название файла

def OT_date(event):
    def print_sel(): # поиск по дате от введенной в главном меню
        textBoxFromDate.delete(0, END)
        DATA = str(cal.selection_get())
        data = DATA[8] + DATA[9] + "." + DATA[5] + DATA[6] + "." + DATA[0] + DATA[1] + DATA[2] + DATA[3]
        textBoxFromDate.insert(0, data)
        okno.update()

    top = tk.Toplevel(okno)
    cal = Calendar(top,
                   font="Arial 14", selectmode='day',
                   cursor="hand1", year=2018, month=2, day=5)
    cal.pack(fill="both", expand=True)
    ttk.Button(top, text="ok", command=print_sel).pack()

def DO_date(event):
    def print_sel(): #поиск по дате до введенной в главном меню
        textBoxToDate.delete(0, END)
        DATA = str(cal.selection_get())
        data = DATA[8] + DATA[9] + "." + DATA[5] + DATA[6] + "." + DATA[0] + DATA[1] + DATA[2] + DATA[3]
        textBoxToDate.insert(0, data)
        okno.update()

    top = tk.Toplevel(okno)
    cal = Calendar(top,
                   font="Arial 14", selectmode='day',
                 cursor="hand1", year=2018, month=2, day=5) #настройка пунктов для даты
    cal.pack(fill="both", expand=True)
    ttk.Button(top, text="ok", command=print_sel).pack()

okno = Tk() # Главное меню
okno.resizable(0, 0) #запрещает изменять размер главного меню
okno.title("Анализ уязвимости")  # Название окна
okno.geometry("520x600")  # Размер окна
okno.configure(background='#3034ab') #цвет фона
labelLow = Label(okno, width=13, height=2, bg='#fabbbf', font='Times 13', text="Низкий") #надпись уровня низкий
labelLow.place(x=0, y=80) #расположение надписи
labelLowOut = Label(okno, bg='#99ffcc', font='Times 15', fg='black', width=5) #поля уровня низкий
labelLowOut.place(x=0, y=125) #расположение поля
labelMid = Label(okno, width=13, height=2, bg='#f78d94', font='Times 13', text="Средний") #надпись уровня средний
labelMid.place(x=0, y=150) #расположение надписи
labelMidOut = Label(okno, bg='#99ffcc', font='Times 15', fg='black', width=5) #поля уровня средний
labelMidOut.place(x=0, y=195) #расположение поля
labelHigh = Label(okno, width=13, height=2, bg='#f55d67', font='Times 13', text="Высокий") #надпись уровня высокий
labelHigh.place(x=0, y=220) #расположение надписи
labelHighOut = Label(okno, bg='#99ffcc', font='Times 15', fg='black', width=5) #поля уровня высокий
labelHighOut.place(x=0, y=265) #расположение поля
labelCrit = Label(okno, width=13, height=2, bg='#f22c39', font='Times 13', text="Критический") #надпись уровня критический
labelCrit.place(x=0, y=295) #расположение надписи
labelCritOut = Label(okno, bg='#99ffcc', font='Times 15', fg='black', width=5) #поля уровня критический
labelCritOut.place(x=0, y=340) #расположение поля
radioButtonDateVar = BooleanVar()  #кнопоки
radioButtonDateVar.set(0) # по умолчанию не выбрано
radioButtonDateOn = Radiobutton(okno, text="По дате", bg='#afdafc', variable=radioButtonDateVar, value=1) #свойства кнопки
radioButtonDateOn.bind('<Button-1>', by_date) #бинд функциии после нажатия, фильтрация по дате
radioButtonDateOn.pack(anchor=W) #растяжения элемента на west(левая сторона)
radioButtonDateOn.place(x=0, y=380) #расположение
radioButtonwithoutDate = Radiobutton(okno, text="За все время", bg='#afdafc', variable=radioButtonDateVar, value=0) #свойства кнопки
radioButtonwithoutDate.bind('<Button-1>', without_date ) #бинд функциии после нажатия, фильтрация без конкретной даты
radioButtonwithoutDate.pack(anchor=W) #растяжения элемента на west(левая сторона)
radioButtonwithoutDate.place(x=0, y=410) #расположение кнопки
combo = Combobox(okno) #список
combo['values'] = ('CentOS', 'Red Hat Enterprise', 'Red Hat Enterprise Linux', 'Red Hat Inc.') #содержание списка
combo.current(0)  # установите вариант по умолчанию CentOS
combo.grid(column=0, row=0) # количество колонок и строк в списке
combo.place(x=230, y=38) #расположение списка
var1 = BooleanVar() # переменная булевская для состояния кнопки 0/1 выбрано/не выбрано
var1.set(0) #не выбрано по дефолту
c1 = Checkbutton(okno, bg = '#afdafc', text="Потенциальная уязвимость", variable=var1, onvalue=1, offvalue=0) #свойства
c1.pack(anchor=W, padx=10) # позиционирование копки. растяжения элемента на west(левая сторона). отступы по горизонтали
c1.grid(column=0, row=0) # количество колонок и строк
c1.place(x=210, y=60) #расположение
buttonobnow = Button(okno, bg='#027ef2', font='Times 12', text="Обновить базу", width=13, height=2) #свойства кнопки
buttonobnow.bind('<Button-1>', donwloade ) # костыль или может и нет
buttonobnow.bind('<Button-1>', AnalysysWithDate ) # после нажатия кнопки выполняется загрузка базы и анализ по дате
buttonobnow.place(x=80, y=515) #расположение кнопки
buttonClear = Button(okno, bg='#027ef2', font='Times 12', text="Удалить", width=13, height=2) #свойства кнопки
buttonClear.place(x=210, y=515) #расположение кнопки
buttonClear.bind('<Button-1>', Clear) #бинд функциии после нажатия, чистануть все поля
buttonSave = Button(okno, bg='#027ef2', font='Times 12', text="Сохранить", width=13, height=2) #свойства кнопки
buttonSave.place(x=340, y=515) #расположение кнопки
buttonSave.bind('<Button-1>', SaveDocx)
buttonDiagram = Button(okno, bg='#7fc7ff', font='Times 12', text="диаграмма", height=2) #свойства кнопки
buttonDiagram.place(x=190, y=460) #расположение кнопки
buttonDiagram.bind('<Button-1>', diagramma) #бинд функциии после нажатия, вызов диаграммы
labelDate = Label(okno, text="Введите дату:", state=DISABLED, bg='#addfad', font='Times 13', fg='#000', width=30) #свойства надписи
labelDate.place(x=175, y=380) #расположение кнопки
labelFromDate = Button(okno, text=" От:", state=DISABLED, bg='#ffd1d1', fg='black', width=5) #свойства кнопки
labelFromDate.place(x=175, y=410) #расположение кнопки
labelFromDate.bind('<Button-1>', OT_date) #бинд функциии после нажатия, определяет от какой даты фильтровать записи
textBoxFromDate=Entry(okno, state=DISABLED, width=15) #свойства поля
textBoxFromDate.place(x=220, y=410) #расположение кнопки
labelToDate = Button(okno, text="До:", state=DISABLED, bg='#ffd1d1', fg='black', width=5) #свойства кнопки
labelToDate.place(x=315, y=410) #расположение кнопки
labelToDate.bind('<Button-1>', DO_date) #бинд функциии после нажатия, определяет до какой даты фильтровать записи
textBoxToDate = Entry(okno, state=DISABLED, width=15) #свойства поля
textBoxToDate.place(x=360, y=410) #расположение кнопки
labelDateInfo = Label(okno, text="Анализ уязвимостей", bg='#78b6f0', font='Times 20', fg='#0f0901', width=50) #свойства надписи
labelDateInfo.pack()
labelToInfo = Label(okno, bg='#78b6f0', fg='black', width=20) #свойства надписи
okno.mainloop()